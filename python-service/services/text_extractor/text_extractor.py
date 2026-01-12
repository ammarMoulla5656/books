"""
Text Extractor - Main Module
Unified interface for extracting text from PDF, DOCX, and ABX files.
"""

import logging
from pathlib import Path
from typing import Optional

from models import ExtractedText, PageContent
from config import settings
from .pdf_extractor import PDFExtractor
from .ocr_processor import OCRProcessor
from .abx_extractor import ABXExtractor

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Unified text extraction from documents.
    Automatically detects file type and chooses appropriate extraction method.
    """

    def __init__(self):
        self.pdf_extractor = PDFExtractor()
        self.abx_extractor = ABXExtractor()
        self.ocr_processor = None  # Lazy loaded

    async def extract(
        self,
        file_path: str,
        use_ocr: bool = False,
        ocr_provider: str = "easyocr"
    ) -> ExtractedText:
        """
        Extract text from a document file.

        Args:
            file_path: Path to the document
            use_ocr: Force OCR even for text-based PDFs
            ocr_provider: OCR provider (easyocr, tesseract, google)

        Returns:
            ExtractedText with all pages and metadata
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = path.suffix.lower()

        if file_ext == ".pdf":
            return await self._extract_from_pdf(file_path, use_ocr, ocr_provider)
        elif file_ext in [".docx", ".doc"]:
            return await self._extract_from_docx(file_path)
        elif file_ext == ".abx":
            return await self.abx_extractor.extract(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    async def _extract_from_pdf(
        self,
        file_path: str,
        use_ocr: bool,
        ocr_provider: str
    ) -> ExtractedText:
        """Extract text from PDF, using OCR if needed."""

        # First try normal text extraction
        extracted, is_scanned = await self.pdf_extractor.extract(file_path)

        # Use OCR if:
        # 1. PDF is scanned (image-based)
        # 2. User explicitly requested OCR
        if is_scanned or use_ocr:
            logger.info("Using OCR for text extraction")

            if self.ocr_processor is None:
                self.ocr_processor = OCRProcessor(provider=ocr_provider)

            ocr_pages = await self.ocr_processor.process_pdf(file_path)

            return ExtractedText(
                text="\n\n".join(p.text for p in ocr_pages),
                pages=ocr_pages,
                total_pages=len(ocr_pages),
                is_scanned=True,
                extraction_method=f"ocr_{ocr_provider}"
            )

        return extracted

    async def _extract_from_docx(self, file_path: str) -> ExtractedText:
        """Extract text from DOCX file."""
        try:
            from docx import Document

            doc = Document(file_path)
            pages = []
            current_text = []
            page_num = 1

            # DOCX doesn't have real page breaks, so we estimate
            # based on paragraph count or explicit page breaks
            for para in doc.paragraphs:
                text = para.text.strip()

                # Check for page break (approximation)
                if self._is_page_break(para) and current_text:
                    pages.append(PageContent(
                        page_number=page_num,
                        text="\n".join(current_text),
                        is_ocr=False
                    ))
                    current_text = []
                    page_num += 1

                if text:
                    current_text.append(text)

            # Add last page
            if current_text:
                pages.append(PageContent(
                    page_number=page_num,
                    text="\n".join(current_text),
                    is_ocr=False
                ))

            full_text = "\n\n".join(p.text for p in pages)

            return ExtractedText(
                text=full_text,
                pages=pages,
                total_pages=len(pages),
                is_scanned=False,
                extraction_method="python-docx"
            )

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _is_page_break(self, paragraph) -> bool:
        """Check if paragraph contains a page break."""
        try:
            from docx.oxml.ns import qn

            # Check for page break in paragraph
            for run in paragraph.runs:
                if run._r.xml.find(qn('w:br')) is not None:
                    br = run._r.find(qn('w:br'))
                    if br is not None and br.get(qn('w:type')) == 'page':
                        return True
            return False
        except Exception:
            return False

    async def get_document_info(self, file_path: str) -> dict:
        """Get document metadata and information."""
        path = Path(file_path)
        file_ext = path.suffix.lower()

        info = {
            "filename": path.name,
            "file_size": path.stat().st_size,
            "file_type": file_ext[1:].upper(),
        }

        if file_ext == ".pdf":
            pdf_info = await self.pdf_extractor.get_pdf_info(file_path)
            info.update(pdf_info)
        elif file_ext in [".docx", ".doc"]:
            docx_info = await self._get_docx_info(file_path)
            info.update(docx_info)
        elif file_ext == ".abx":
            abx_info = await self._get_abx_info(file_path)
            info.update(abx_info)

        return info

    async def _get_docx_info(self, file_path: str) -> dict:
        """Get DOCX metadata."""
        try:
            from docx import Document

            doc = Document(file_path)
            core_props = doc.core_properties

            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "page_count": len(doc.paragraphs) // 30 + 1,  # Rough estimate
                "has_toc": False  # Will need to detect from content
            }
        except Exception as e:
            logger.error(f"Failed to get DOCX info: {e}")
            return {}

    async def _get_abx_info(self, file_path: str) -> dict:
        """Get ABX metadata."""
        try:
            # Extract content to get basic info
            extracted = await self.abx_extractor.extract(file_path)

            return {
                "title": "",  # Could be extracted from metadata
                "author": "",  # Could be extracted from metadata
                "subject": "",
                "page_count": extracted.total_pages,
                "has_toc": False  # Will need to detect from content
            }
        except Exception as e:
            logger.error(f"Failed to get ABX info: {e}")
            return {}
